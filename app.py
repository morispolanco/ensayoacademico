import streamlit as st
import requests
import json
from docx import Document
from docx.shared import Inches
from docx.enum.dml import MSO_THEME_COLOR_INDEX
from io import BytesIO
import re

# Configuración de la página
st.set_page_config(page_title="Generador de Ensayos Académicos", page_icon="📝", layout="wide")

# Función para crear la columna de información
def crear_columna_info():
    st.markdown("""
    ## Sobre esta aplicación

    Esta aplicación es un Generador de Ensayos Académicos que utiliza inteligencia artificial para ayudarte a crear ensayos estructurados y bien documentados.

    ### Cómo usar la aplicación:

    1. Ingrese el título de su ensayo.
    2. Escriba la tesis que quiere defender.
    3. Ingrese los autores que desea citar (separados por comas).
    4. Haga clic en "Generar ensayo académico" para obtener el resultado.
    5. Lea el ensayo generado con enlaces a las fuentes citadas.
    6. Si lo desea, descargue un documento DOCX con toda la información y enlaces activos.

    ### Autor y actualización:
    **Moris Polanco**, 27 ag 2024

    ### Cómo citar esta aplicación (formato APA):
    Polanco, M. (2024). *Generador de Ensayos Académicos* [Aplicación web]. https://ensayoacademico.streamlit.app

    ---
    **Nota:** Esta aplicación utiliza inteligencia artificial para generar contenido basado en información disponible en línea. Siempre verifique la información con fuentes académicas para un análisis más profundo.
    """)

# Título de la aplicación
st.title("Generador de Ensayos Académicos")

# Crear un diseño de dos columnas
col1, col2 = st.columns([1, 2])

# Columna de información
with col1:
    crear_columna_info()

# Columna principal
with col2:
    # Acceder a las claves de API de los secretos de Streamlit
    TOGETHER_API_KEY = st.secrets["TOGETHER_API_KEY"]
    SERPLY_API_KEY = st.secrets["SERPLY_API_KEY"]

    def buscar_informacion(query):
        url = "https://api.serply.io/v1/scholar"
        params = {
            "q": query
        }
        headers = {
            'X-Api-Key': SERPLY_API_KEY,
            'Content-Type': 'application/json'
        }
        response = requests.get(url, headers=headers, params=params)
        return response.json()

    def generar_ensayo(titulo, tesis, autores, fuentes):
        url = "https://api.together.xyz/inference"
        fuentes_str = "\n".join([f"- {fuente}" for fuente in fuentes])
        prompt = f"""Escribe un ensayo académico con el título "{titulo}". 
        La tesis principal a defender es: "{tesis}"
        Incluye citas y discusiones de los siguientes autores: {autores}.
        Utiliza y cita las siguientes fuentes en tu ensayo:
        {fuentes_str}
        El ensayo debe tener la siguiente estructura:
        1. Introducción (presenta el tema y la tesis)
        2. Desarrollo (argumenta la tesis, utilizando las fuentes y autores proporcionados)
        3. Conclusión (resume los puntos principales y reafirma la tesis)

        No escribas las instrucciones ni ejemplos ni las referencias al principio.
        Asegúrate de incluir citas en el texto y una lista de referencias al final. 
        Para cada cita en el texto, usa el formato [Autor, Año] y asegúrate de que corresponda con una entrada en la lista de referencias."""

        payload = json.dumps({
            "model": "mistralai/Mixtral-8x7B-Instruct-v0.1",
            "prompt": prompt,
            "max_tokens": 3048,
            "temperature": 0.7,
            "top_p": 0.7,
            "top_k": 50,
            "repetition_penalty": 1,
            "stop": ["Título:"]
        })
        headers = {
            'Authorization': f'Bearer {TOGETHER_API_KEY}',
            'Content-Type': 'application/json'
        }
        response = requests.post(url, headers=headers, data=payload)
        return response.json()['output']['choices'][0]['text'].strip()

    def add_hyperlink(paragraph, url, text):
        # This gets access to the document.xml.rels file and gets a new relation id value
        part = paragraph.part
        r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

        # Create the w:hyperlink tag and add needed values
        hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
        hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

        # Create a w:r element
        new_run = docx.oxml.shared.OxmlElement('w:r')

        # Create a new w:rPr element
        rPr = docx.oxml.shared.OxmlElement('w:rPr')

        # Add color if it is needed
        c = docx.oxml.shared.OxmlElement('w:color')
        c.set(docx.oxml.shared.qn('w:val'), '0000FF')
        rPr.append(c)

        # Add underlining
        u = docx.oxml.shared.OxmlElement('w:u')
        u.set(docx.oxml.shared.qn('w:val'), 'single')
        rPr.append(u)

        # Join all the xml elements together add add the required text to the w:r element
        new_run.append(rPr)
        new_run.text = text
        hyperlink.append(new_run)

        paragraph._p.append(hyperlink)

        return hyperlink

    def create_docx(titulo, contenido, fuentes):
        doc = Document()
        doc.add_heading('Ensayo Académico', 0)

        doc.add_heading(titulo, level=1)
        
        # Dividir el contenido en párrafos
        parrafos = contenido.split('\n\n')
        
        for parrafo in parrafos:
            p = doc.add_paragraph()
            # Buscar citas en el formato [Autor, Año]
            citas = re.findall(r'\[([^\]]+)\]', parrafo)
            partes = re.split(r'\[([^\]]+)\]', parrafo)
            
            for i, parte in enumerate(partes):
                if i % 2 == 0:  # Texto normal
                    p.add_run(parte)
                else:  # Cita
                    # Buscar la fuente correspondiente
                    for fuente in fuentes:
                        if parte.lower() in fuente.lower():
                            # Extraer el enlace de la fuente
                            enlace = fuente.split(': ')[-1]
                            # Añadir el hipervínculo
                            add_hyperlink(p, enlace, f'[{parte}]')
                            break
                    else:
                        # Si no se encuentra una fuente correspondiente, añadir el texto sin hipervínculo
                        p.add_run(f'[{parte}]')

        doc.add_paragraph('\nNota: Este documento fue generado por un asistente de IA. Verifica la información con fuentes académicas para un análisis más profundo.')

        return doc

    # Interfaz de usuario
    titulo = st.text_input("Ingrese el título de su ensayo:")
    tesis = st.text_area("Escriba la tesis que quiere defender:")
    autores = st.text_input("Ingrese los autores que desea citar (separados por comas):")

    if st.button("Generar ensayo académico"):
        if titulo and tesis and autores:
            with st.spinner("Buscando información y generando ensayo..."):
                # Buscar información relevante para cada autor
                autores_lista = [autor.strip() for autor in autores.split(',')]
                fuentes = []
                for autor in autores_lista:
                    resultados_busqueda = buscar_informacion(f"{autor} {titulo}")
                    fuentes.extend([f"{resultado['title']}: {resultado['link']}" for resultado in resultados_busqueda.get('results', [])[:2]])
                
                # Generar ensayo
                ensayo = generar_ensayo(titulo, tesis, autores, fuentes)

                # Mostrar ensayo
                st.write("Ensayo generado:")
                
                # Dividir el ensayo en párrafos
                parrafos = ensayo.split('\n\n')
                
                for parrafo in parrafos:
                    # Buscar citas en el formato [Autor, Año]
                    citas = re.findall(r'\[([^\]]+)\]', parrafo)
                    partes = re.split(r'\[([^\]]+)\]', parrafo)
                    
                    nuevo_parrafo = ""
                    for i, parte in enumerate(partes):
                        if i % 2 == 0:  # Texto normal
                            nuevo_parrafo += parte
                        else:  # Cita
                            # Buscar la fuente correspondiente
                            for fuente in fuentes:
                                if parte.lower() in fuente.lower():
                                    # Extraer el enlace de la fuente
                                    enlace = fuente.split(': ')[-1]
                                    # Añadir el hipervínculo
                                    nuevo_parrafo += f'[{parte}]({enlace})'
                                    break
                            else:
                                # Si no se encuentra una fuente correspondiente, dejar el texto sin hipervínculo
                                nuevo_parrafo += f'[{parte}]'
                    
                    st.markdown(nuevo_parrafo)
                    st.write("")  # Añadir un espacio entre párrafos

                # Crear documento DOCX
                doc = create_docx(titulo, ensayo, fuentes)

                # Guardar el documento DOCX en memoria
                docx_file = BytesIO()
                doc.save(docx_file)
                docx_file.seek(0)

                # Opción para exportar a DOCX
                st.download_button(
                    label="Descargar ensayo como DOCX",
                    data=docx_file,
                    file_name=f"{titulo.lower().replace(' ', '_')}_ensayo.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )

        else:
            st.warning("Por favor, complete todos los campos antes de generar el ensayo.")
